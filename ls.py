import argparse
import logging
import sys
from pathlib import Path
from fnmatch import fnmatch

import docx
from docx.shared import Pt
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

try:
    from progress.bar import IncrementalBar
except Exception:
    IncrementalBar = None

try:
    from colorama import init as colorama_init, Fore, Style
except Exception:
    colorama_init = None
    Fore = Style = None


# ===== Настройки =====
MAX_DOC_CHARS = 800_000  # лимит символов кода на один docx (по тексту файлов)

# ГОСТ 19.106-78
CODE_FONT_NAME = "Courier New"
CODE_FONT_SIZE_PT = 12

# Колонки таблицы (мм)
NUM_COL_WIDTH_MM = 16
CODE_COL_WIDTH_MM = 170

APP_NAME = "GOST Listing Generator"


# ===== Цветные логи без времени/уровня =====
class _ColorFormatter(logging.Formatter):
    def format(self, record: logging.LogRecord) -> str:
        msg = record.getMessage()

        if Fore is None or Style is None:
            return msg

        if record.levelno >= logging.ERROR:
            return f"{Fore.RED}{msg}{Style.RESET_ALL}"
        if record.levelno >= logging.WARNING:
            return f"{Fore.YELLOW}{msg}{Style.RESET_ALL}"
        # INFO
        return f"{Fore.CYAN}{msg}{Style.RESET_ALL}"


def setup_logging() -> logging.Logger:
    if colorama_init is not None:
        colorama_init(autoreset=True)

    logger = logging.getLogger(APP_NAME)
    logger.setLevel(logging.INFO)
    logger.handlers.clear()

    # ВАЖНО: логи в stderr, чтобы не ломать прогресс-бар (stdout)
    h = logging.StreamHandler(sys.stderr)
    h.setLevel(logging.INFO)
    h.setFormatter(_ColorFormatter("%(message)s"))
    logger.addHandler(h)
    return logger


def c_info(msg: str) -> str:
    if Fore is None or Style is None:
        return msg
    return f"{Fore.CYAN}{msg}{Style.RESET_ALL}"


def c_ok(msg: str) -> str:
    if Fore is None or Style is None:
        return msg
    return f"{Fore.GREEN}{msg}{Style.RESET_ALL}"


def c_warn(msg: str) -> str:
    if Fore is None or Style is None:
        return msg
    return f"{Fore.YELLOW}{msg}{Style.RESET_ALL}"


def c_err(msg: str) -> str:
    if Fore is None or Style is None:
        return msg
    return f"{Fore.RED}{msg}{Style.RESET_ALL}"


# ===== База =====
def app_dir() -> Path:
    """
    В .py: папка со скриптом.
    В .exe (PyInstaller): папка, откуда запускают exe.
    """
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def xml_safe_text(s: str) -> str:
    out = []
    for ch in s:
        code = ord(ch)
        if code in (0x9, 0xA, 0xD):
            out.append(ch)
        elif 0x20 <= code <= 0xD7FF:
            out.append(ch)
        elif 0xE000 <= code <= 0xFFFD:
            out.append(ch)
        elif 0x10000 <= code <= 0x10FFFF:
            out.append(ch)
        else:
            continue
    return "".join(out)


# ===== Первичная инициализация окружения =====
IGNORE_TEMPLATE = """# ignore.txt — маски исключений (fnmatch)
# Строки с # — комментарии.
# Примеры:
# __pycache__/
# .git/
# .venv/
# node_modules/
# *.log
# *.tmp
"""

TARGETS_README = """Сюда положи проекты (каждый проект — отдельная папка).

Пример:
targets/
  ProjectA/
    main.py
    src/...
  ProjectB/
    ...

Затем запусти программу ещё раз — она соберёт листинги в listing_out/.
"""


def ensure_first_run_layout(base: Path, logger: logging.Logger) -> bool:
    """
    True => это первый запуск, структура создана, нужно завершиться.
    False => targets уже есть, можно выполнять обработку.
    """
    targets = base / "targets"
    listing_out = base / "listing_out"
    ignore_txt = base / "ignore.txt"

    if targets.exists():
        return False

    logger.info(c_info("Первый запуск: создаю рабочую структуру..."))
    targets.mkdir(parents=True, exist_ok=True)
    listing_out.mkdir(parents=True, exist_ok=True)

    if not ignore_txt.exists():
        ignore_txt.write_text(IGNORE_TEMPLATE, encoding="utf-8")

    readme = targets / "README.txt"
    if not readme.exists():
        readme.write_text(TARGETS_README, encoding="utf-8")

    logger.info(c_ok("Создано: targets/, listing_out/, ignore.txt"))
    logger.info(c_info("Положи проекты в targets/ и запусти программу ещё раз."))
    return True


# ===== Ignore (ignore.txt рядом с программой или в targets/) =====
def normalize_patterns(lines: list[str]) -> list[str]:
    patterns: list[str] = []
    for raw in lines:
        s = raw.strip()
        if not s or s.startswith("#"):
            continue

        # удобство: ".env" -> "*.env" (если это не путь и без wildcard)
        if s.startswith(".") and ("/" not in s) and ("*" not in s) and ("?" not in s) and ("[" not in s):
            s = f"*{s}"

        s = s.replace("\\", "/")
        patterns.append(s)
    return patterns


def load_ignore_patterns_auto(base: Path) -> list[str]:
    candidates = [
        base / "ignore.txt",
        base / "targets" / "ignore.txt",
    ]
    for c in candidates:
        if c.exists() and c.is_file():
            return normalize_patterns(c.read_text(encoding="utf-8", errors="ignore").splitlines())
    return []


def is_ignored(rel_posix: str, name: str, patterns: list[str]) -> bool:
    for pat in patterns:
        # директория: 'build/'
        if pat.endswith("/"):
            d = pat[:-1]
            if rel_posix == d or rel_posix.startswith(d + "/"):
                return True
            if f"/{d}/" in f"/{rel_posix}/":
                return True
            continue

        anchored = pat.startswith("/")
        p = pat[1:] if anchored else pat

        if anchored:
            if fnmatch(rel_posix, p):
                return True
        else:
            if fnmatch(rel_posix, p) or fnmatch(name, p):
                return True
    return False


def iter_project_files(project_dir: Path, out_root: Path, patterns: list[str]) -> list[Path]:
    files: list[Path] = []
    for p in project_dir.rglob("*"):
        if not p.is_file():
            continue

        # если вдруг listing_out внутри проекта
        try:
            p.relative_to(out_root)
            continue
        except ValueError:
            pass

        rel = p.relative_to(project_dir).as_posix()
        if is_ignored(rel, p.name, patterns):
            continue

        files.append(p)

    return sorted(files, key=lambda x: x.as_posix().lower())


# ===== DOCX: стили и заголовки =====
def ensure_code_style(doc: docx.Document) -> None:
    styles = doc.styles
    if "CodeStyle" not in [s.name for s in styles]:
        st = styles.add_style("CodeStyle", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = CODE_FONT_NAME
        st.font.size = Pt(CODE_FONT_SIZE_PT)
        st.font.bold = False
        st.font.italic = False
        st.font.underline = False


def add_header(doc: docx.Document, appendix_label: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"Приложение {appendix_label}")
    r.font.name = CODE_FONT_NAME
    r.font.size = Pt(CODE_FONT_SIZE_PT)
    r.bold = False
    r.italic = False
    r.underline = False

    p.alignment = 1  # center
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = 0
    pf.left_indent = 0
    pf.right_indent = 0
    pf.line_spacing = 1.0


def add_listing_heading(doc: docx.Document, listing_no: int, rel_name: str, part_suffix: str = ""):
    title = f"Листинг {listing_no} — {rel_name}"
    if part_suffix:
        title += f" ({part_suffix})"

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.name = CODE_FONT_NAME
    r.font.size = Pt(CODE_FONT_SIZE_PT)
    r.bold = False
    r.italic = False
    r.underline = False

    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = 0
    pf.left_indent = 0
    pf.right_indent = 0
    pf.line_spacing = 1.0
    pf.keep_with_next = True

    return p._p


def add_separator_paragraph(doc: docx.Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run("")
    r.font.name = CODE_FONT_NAME
    r.font.size = Pt(CODE_FONT_SIZE_PT)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def split_keep_all_lines(text: str) -> list[str]:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text.split("\n")


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def make_project_out_dir(listing_out: Path, project_name: str) -> Path:
    p = listing_out / project_name
    p.mkdir(parents=True, exist_ok=True)
    return p


def doc_name(project_out: Path, project_name: str, doc_idx: int) -> Path:
    return project_out / f"{project_name}_listing_{doc_idx}.docx"


def new_doc(appendix_label: str) -> docx.Document:
    d = docx.Document()
    ensure_code_style(d)
    add_header(d, appendix_label)
    return d


def index_to_label(index: int, alphabet: list[str]) -> str:
    base = len(alphabet)
    n = index
    chars: list[str] = []
    while True:
        chars.append(alphabet[n % base])
        n = n // base - 1
        if n < 0:
            break
    return "".join(reversed(chars))


# ===== Таблица: только внешний контур =====
def _dxa_from_mm(mm: int) -> str:
    return str(int(mm * 56.7))


def _make_tcW(mm: int) -> OxmlElement:
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), _dxa_from_mm(mm))
    return tcW


def _set_tc_nowrap(tcPr: OxmlElement) -> None:
    noWrap = OxmlElement("w:noWrap")
    tcPr.append(noWrap)


def _make_pPr_gost() -> OxmlElement:
    pPr = OxmlElement("w:pPr")

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    pPr.append(jc)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:line"), "240")
    pPr.append(spacing)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    pPr.append(ind)

    suppressHyphens = OxmlElement("w:suppressAutoHyphens")
    suppressHyphens.set(qn("w:val"), "1")
    pPr.append(suppressHyphens)

    return pPr


def _make_rPr_gost() -> OxmlElement:
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), CODE_FONT_NAME)
    rFonts.set(qn("w:hAnsi"), CODE_FONT_NAME)
    rFonts.set(qn("w:cs"), CODE_FONT_NAME)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(CODE_FONT_SIZE_PT * 2))
    rPr.append(sz)

    return rPr


_PPR_XML = None
_RPR_XML = None


def _ppr_xml() -> str:
    global _PPR_XML
    if _PPR_XML is None:
        _PPR_XML = _make_pPr_gost().xml
    return _PPR_XML


def _rpr_xml() -> str:
    global _RPR_XML
    if _RPR_XML is None:
        _RPR_XML = _make_rPr_gost().xml
    return _RPR_XML


def _clone(xml: str) -> OxmlElement:
    return parse_xml(xml)


def _make_cell_paragraph(text: str) -> OxmlElement:
    text = xml_safe_text(text)

    p = OxmlElement("w:p")
    p.append(_clone(_ppr_xml()))

    r = OxmlElement("w:r")
    r.append(_clone(_rpr_xml()))

    t = OxmlElement("w:t")
    if text.startswith(" ") or "  " in text:
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)

    p.append(r)
    return p


def _set_outer_borders_only(tblPr: OxmlElement) -> None:
    tblBorders = OxmlElement("w:tblBorders")

    def border(tag: str, val: str) -> OxmlElement:
        b = OxmlElement(tag)
        b.set(qn("w:val"), val)
        if val != "nil":
            b.set(qn("w:sz"), "4")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")
        return b

    tblBorders.append(border("w:top", "single"))
    tblBorders.append(border("w:left", "single"))
    tblBorders.append(border("w:bottom", "single"))
    tblBorders.append(border("w:right", "single"))
    tblBorders.append(border("w:insideH", "nil"))
    tblBorders.append(border("w:insideV", "nil"))

    tblPr.append(tblBorders)


def build_code_table_xml(lines: list[str], start_line_no: int = 1) -> OxmlElement:
    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    _set_outer_borders_only(tblPr)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "auto")
    tblW.set(qn("w:w"), "0")
    tblPr.append(tblW)

    tbl.append(tblPr)

    tblGrid = OxmlElement("w:tblGrid")
    gridCol1 = OxmlElement("w:gridCol")
    gridCol1.set(qn("w:w"), _dxa_from_mm(NUM_COL_WIDTH_MM))
    gridCol2 = OxmlElement("w:gridCol")
    gridCol2.set(qn("w:w"), _dxa_from_mm(CODE_COL_WIDTH_MM))
    tblGrid.append(gridCol1)
    tblGrid.append(gridCol2)
    tbl.append(tblGrid)

    for i, line in enumerate(lines):
        safe_line = xml_safe_text(line)

        tr = OxmlElement("w:tr")

        tc1 = OxmlElement("w:tc")
        tcPr1 = OxmlElement("w:tcPr")
        tcPr1.append(_make_tcW(NUM_COL_WIDTH_MM))
        _set_tc_nowrap(tcPr1)
        tc1.append(tcPr1)
        tc1.append(_make_cell_paragraph(str(start_line_no + i)))

        tc2 = OxmlElement("w:tc")
        tcPr2 = OxmlElement("w:tcPr")
        tcPr2.append(_make_tcW(CODE_COL_WIDTH_MM))
        _set_tc_nowrap(tcPr2)
        tc2.append(tcPr2)
        tc2.append(_make_cell_paragraph(safe_line))

        tr.append(tc1)
        tr.append(tc2)
        tbl.append(tr)

    return tbl


def add_code_table_for_lines_fast_after(heading_p, lines: list[str], start_line_no: int = 1) -> None:
    tbl = build_code_table_xml(lines, start_line_no=start_line_no)
    heading_p.addnext(tbl)


# ===== Разбиение больших файлов по строкам под лимит docx =====
def compute_line_blocks_by_char_limit(lines: list[str], max_chars: int) -> list[tuple[int, int]]:
    blocks: list[tuple[int, int]] = []
    n = len(lines)
    i = 0
    while i < n:
        j = i
        acc = 0
        while j < n:
            add_len = len(lines[j]) + 1
            if acc > 0 and (acc + add_len > max_chars):
                break
            if acc == 0 and add_len > max_chars:
                j += 1
                break
            acc += add_len
            j += 1
        blocks.append((i, j))
        i = j
    return blocks


# ===== Основная обработка проекта =====
def process_project(project_dir: Path, listing_out: Path, patterns: list[str], appendix_label: str, logger: logging.Logger) -> None:
    project_name = project_dir.name
    project_out = make_project_out_dir(listing_out, project_name)

    files = iter_project_files(project_dir, listing_out, patterns)
    if not files:
        logger.warning(c_warn(f"[{project_name}] Нет файлов для обработки (пусто или всё отфильтровано)."))
        return

    logger.info(c_info(f"[{project_name}] Файлов к обработке: {len(files)}"))

    bar = None
    if IncrementalBar is not None and sys.stdout.isatty():
        try:
            bar = IncrementalBar(
                f"{project_name} ({appendix_label})",
                max=len(files),
                suffix=" %(index).d/%(max).d - %(percent).1f%% - %(elapsed).ds",
                file=sys.stdout,  # важно: бар в stdout
            )
        except TypeError:
            # если версия progress не поддерживает file=
            bar = IncrementalBar(
                f"{project_name} ({appendix_label})",
                max=len(files),
                suffix=" %(index).d/%(max).d - %(percent).1f%% - %(elapsed).ds",
            )

    doc_idx = 1
    listing_no = 1
    doc = new_doc(appendix_label)
    doc_chars = 0

    for f in files:
        rel = f.relative_to(project_dir).as_posix()
        try:
            content = read_text(f)
        except (OSError, PermissionError):
            logger.warning(c_warn(f"[{project_name}] Пропуск (нет доступа): {rel}"))
            if bar:
                bar.next()
            continue

        content = xml_safe_text(content)
        content_len = len(content)

        if content_len <= MAX_DOC_CHARS:
            if doc_chars > 0 and (doc_chars + content_len > MAX_DOC_CHARS):
                out_path = doc_name(project_out, project_name, doc_idx)
                doc.save(str(out_path))

                doc_idx += 1
                doc = new_doc(appendix_label)
                doc_chars = 0

            heading_p = add_listing_heading(doc, listing_no, rel)
            lines = split_keep_all_lines(content)
            add_code_table_for_lines_fast_after(heading_p, lines, start_line_no=1)
            add_separator_paragraph(doc)

            listing_no += 1
            doc_chars += content_len

        else:
            if doc_chars > 0:
                out_path = doc_name(project_out, project_name, doc_idx)
                doc.save(str(out_path))
                doc_idx += 1
                doc = new_doc(appendix_label)
                doc_chars = 0

            all_lines = split_keep_all_lines(content)
            blocks = compute_line_blocks_by_char_limit(all_lines, MAX_DOC_CHARS)
            total_parts = len(blocks)

            logger.info(c_info(f"[{project_name}] Большой файл: {rel} -> частей: {total_parts}"))

            for part_idx, (a, b) in enumerate(blocks, start=1):
                heading_p = add_listing_heading(
                    doc, listing_no, rel, part_suffix=f"часть {part_idx}/{total_parts}"
                )
                add_code_table_for_lines_fast_after(heading_p, all_lines[a:b], start_line_no=a + 1)
                add_separator_paragraph(doc)

                out_path = doc_name(project_out, project_name, doc_idx)
                doc.save(str(out_path))
                doc_idx += 1

                if part_idx != total_parts:
                    doc = new_doc(appendix_label)
                doc_chars = 0

            listing_no += 1
            doc = new_doc(appendix_label)
            doc_chars = 0

        if bar:
            bar.next()

    if bar:
        bar.finish()

    if doc_chars > 0:
        out_path = doc_name(project_out, project_name, doc_idx)
        doc.save(str(out_path))

    logger.info(c_ok(f"[{project_name}] Готово."))

def pause_if_double_click():
    """
    Пауза только при запуске .exe двойным кликом.
    В консоли (cmd / powershell) паузы не будет.
    """
    if getattr(sys, "frozen", False) and sys.stdout is not None:
        try:
            if not sys.stdin.isatty():
                input("\nНажмите Enter для выхода...")
        except Exception:
            pass
# ===== main =====
def main() -> int:
    logger = setup_logging()

    parser = argparse.ArgumentParser(
        description="ГОСТ-листинги из targets/* в listing_out/* (первый запуск создаёт структуру)."
    )
    parser.parse_args()

    base = app_dir()

    # первый запуск — создать структуру и выйти
    if ensure_first_run_layout(base, logger):
        return 0

    targets = base / "targets"
    listing_out = base / "listing_out"
    listing_out.mkdir(parents=True, exist_ok=True)

    patterns = load_ignore_patterns_auto(base)

    logger.info(c_info(f"База: {base}"))
    logger.info(c_info(f"Папка проектов: {targets}"))
    logger.info(c_info(f"Вывод: {listing_out}"))
    logger.info(c_info(f"Правил ignore: {len(patterns)}"))

    if not targets.exists() or not targets.is_dir():
        logger.error(c_err(f"Нет папки targets: {targets}"))
        return 2

    projects = sorted([p for p in targets.iterdir() if p.is_dir()], key=lambda p: p.name.lower())
    if not projects:
        logger.warning(c_warn("В targets нет проектов (подпапок). Добавь папки проектов и запусти снова."))
        pause_if_double_click()
        return 0

    alphabet = list("АБВГДЕЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ")

    logger.info(c_info(f"Проектов найдено: {len(projects)}"))
    for idx, project in enumerate(projects):
        appendix_label = index_to_label(idx, alphabet)
        logger.info(c_info(f"=== Старт проекта: {project.name} | Приложение: {appendix_label} ==="))
        process_project(project, listing_out, patterns, appendix_label, logger)

    logger.info(c_ok(f"Готово. Результаты: {listing_out}"))

    pause_if_double_click()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
