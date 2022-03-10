import os
from argparse import ArgumentParser
from docx.shared import Pt,Mm
import docx
from progress.bar import IncrementalBar
import math

it=0

j=0
i=1
class List_file:
    def __init__(self,path,ignore,hr,file_name,intcount):
        self.path=path
        self.ignore=ignore
        self.filelist=list()
        self.newlist = list()
        self.hr=hr
        self.file_name=file_name
        self.intcount=intcount

    def list_files(self):
        if os.path.exists(self.path):
            Doc_work.create_doc(self.hr,self.file_name)
            for root, dirs, files in os.walk(self.path):
                for file in files:
                    self.filelist.append(os.path.join(root, file))
            List_file.check_ignore(self)
        else:
            print(f'Каталога \'{self.path}\' не существует')

    def check_ignore(self):
        ignore = self.ignore.split()
        for i in self.filelist:
            found = False
            for j in ignore:
                if f'\\{j}\\' in i:
                    found = j
                    break
            if found:
                continue
            else:
                self.newlist.append(i)
        self.filelist=self.newlist
        self.intcount = math.ceil(len(self.filelist) / int(self.intcount))
        List_file.crutch(self)


    def crutch(self):
        it = len(self.filelist)
        bar = IncrementalBar('Loading...', max=it, suffix=f' %(index).d/%(max).d - %(percent).1f%% - %(elapsed).ds')
        while it > 0:
            try:
                it = it - 1
                entry(path, filelist, hr, file_name, count, bar, intcount)
            except:
                continue
        bar.finish()

    def entry(self):
        global j
        doc = docx.Document(f'{self.file_name} - {Doc_work.o}.docx')
        if len(self.filelist) == 0:
            return True
        for name in self.filelist:
            if j >= int(self.intcount):
                Doc_work.create_doc1(self.file_name)
            else:
                filelist.remove(name)
                j += 1
                name1 = name.replace(f'{path}\\', "")
                name = name.replace("\\", "\\\\")
                dock_formation(doc, name1, name, hr, bar)
                doc.save(f'{file_name} - {o}.docx')
                bar.next()


class Doc_work:
    o=1
    @staticmethod
    def create_doc(hr,file_name):
        doc = docx.Document()
        par = doc.add_paragraph()
        par.add_run(f'Приложение {hr}').bold = True
        par.alignment = 1
        doc.save(f'{file_name} - {Doc_work.o}.docx')

    @staticmethod
    def create_doc1(file_name):
        Doc_work.o += 1
        doc = docx.Document()
        doc.save(f'{file_name} - {Doc_work.o}.docx')
        List_file().entry()

    @classmethod
    def dock_formation(doc, name1, name, hr, bar):
        global i
        p = doc.add_paragraph(f'Листинг {hr}.{i} - {name1}')
        try:
            i += 1
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            info = open(name, encoding="utf8", errors='ignore').read().strip()
            table.cell(0, 0).text = info
            fmt = p.paragraph_format
            fmt.space_before = Mm(3)
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(14)
        except:
            delete_paragraph(p)
            i -= 1
            allTables = doc.tables
            for activeTable in allTables:
                if activeTable.cell(0, 0).paragraphs[0].text == '':
                    activeTable._element.getparent().remove(activeTable._element)

    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None











if __name__ == '__main__':
    parser = ArgumentParser(description="Formation of the program listing")
    parser.add_argument("-td", "--tirgert_dir", dest="path", required=True,
                        help="Directory path")
    parser.add_argument("-ig", "--ignore_dir", dest="ignore", default=".git",
                        help="Ignore directory")
    parser.add_argument("-hr", "--header", dest="hr", default="А",
                        help="Application number")
    parser.add_argument("-o", "--output", dest="file_name", required=True,
                        help="File name")
    parser.add_argument("-n", "--num", dest="intcount", required=True,
                        help="Number of files to split into")
    args = parser.parse_args()

    list_files(args.path, args.ignore, args.hr, args.file_name,args.intcount)



